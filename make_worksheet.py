"""
Generate a scaffolded worksheet (.docx) for Ex 1K – Applications of Simultaneous Equations
55-minute lesson  |  Student-facing
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)

# ── Colour helpers ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x23, 0x5C)
TEAL   = RGBColor(0x00, 0x97, 0x9C)
GOLD   = RGBColor(0xFF, 0xC0, 0x2E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x2E, 0x86, 0x48)
ORANGE = RGBColor(0xE8, 0x57, 0x1A)
PURPLE = RGBColor(0x6A, 0x3D, 0x9A)
RED    = RGBColor(0xC0, 0x20, 0x20)

def set_cell_bg(cell, rgb: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=NAVY, size=18, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def body(doc, text, size=11, color=DGRAY, bold=False, italic=False,
         space_before=2, space_after=2, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def blank_lines(doc, n=3, label="Working space"):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run("_" * 90)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def answer_box(doc, rows=2):
    """A shaded answer box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, LGRAY)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    run = p.add_run("Answer: " + "_" * 60)
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    for _ in range(rows - 1):
        p2 = cell.add_paragraph("_" * 72)
        p2.runs[0].font.size = Pt(11)
        p2.runs[0].font.color.rgb = NAVY
    doc.add_paragraph()

def section_banner(doc, text, bg=NAVY, fg=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = fg
    doc.add_paragraph()

def step_scaffold_table(doc, steps):
    """4-column scaffold for the 4-step method."""
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    colors = [TEAL, GREEN, ORANGE, PURPLE]
    for i, (step_title, step_hint) in enumerate(steps):
        hdr_cell = tbl.cell(0, i)
        set_cell_bg(hdr_cell, colors[i])
        p = hdr_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step_title)
        run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE

        body_cell = tbl.cell(1, i)
        set_cell_bg(body_cell, LGRAY)
        p2 = body_cell.paragraphs[0]
        run2 = p2.add_run(step_hint)
        run2.font.size = Pt(9); run2.font.color.rgb = DGRAY; run2.italic = True
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# WORKSHEET HEADER
# ══════════════════════════════════════════════════════════════════════════════
# Title banner
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
set_cell_bg(cell, NAVY)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Ex 1K  —  Applications of Simultaneous Linear Equations")
r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = WHITE
doc.add_paragraph()

# Student info row
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, lbl in enumerate(["Name: ___________________________",
                          "Date: ____________",
                          "Class: __________",
                          "Score: _____ / 50 XP"]):
    c = tbl2.cell(0, i)
    rn = c.paragraphs[0].add_run(lbl)
    rn.font.size = Pt(10); rn.font.color.rgb = DGRAY
doc.add_paragraph()

# Learning intentions box
tbl3 = doc.add_table(rows=1, cols=1)
cell3 = tbl3.cell(0, 0)
set_cell_bg(cell3, RGBColor(0xE8, 0xF4, 0xFD))
p3 = cell3.paragraphs[0]
r3 = p3.add_run("🎯  Learning Intentions")
r3.font.bold = True; r3.font.size = Pt(12); r3.font.color.rgb = NAVY
for intent in [
    "  ✔  I can define variables clearly from a worded problem.",
    "  ✔  I can write two linear equations from given information.",
    "  ✔  I can solve simultaneous equations and interpret the answer in context.",
    "  ✔  I can identify parallel and coincident lines from their equations.",
]:
    pp = cell3.add_paragraph(intent)
    pp.runs[0].font.size = Pt(10); pp.runs[0].font.color.rgb = DGRAY
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 – WARM UP
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, "⚡  PART 1  |  Warm-Up  (5 minutes)", bg=TEAL)

body(doc, "At the school canteen:", bold=True, size=11, color=NAVY)
body(doc, "   •  2 pies + 3 drinks = $13.00", size=11, color=DGRAY)
body(doc, "   •  4 pies + 1 drink  = $15.00", size=11, color=DGRAY)
body(doc, "Using the 4-step method below, find the price of one pie and one drink.",
     size=11, bold=True, color=NAVY, space_before=6)

step_scaffold_table(doc, [
    ("STEP 1\nDefine Variables",
     'Let p = price of a pie\nLet d = price of a drink'),
    ("STEP 2\nForm Equations",
     'Write equation (1): _______________\nWrite equation (2): _______________'),
    ("STEP 3\nSolve",
     'Show your working below.\nUse elimination or substitution.'),
    ("STEP 4\nAnswer in Context",
     'Write a sentence:\n"The price of one pie is …"'),
])

blank_lines(doc, 5)
answer_box(doc, 2)

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 – THE 4-STEP METHOD (REMINDER)
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, "📐  The 4-Step Method  —  Quick Reference", bg=NAVY)

steps_ref = [
    ("1. Define Variables",
     'Choose a meaningful letter for each unknown. Write e.g. "Let x = …"'),
    ("2. Form Equations",
     "Each piece of information gives ONE equation. You need exactly two equations for two unknowns."),
    ("3. Solve Simultaneously",
     "Use substitution or elimination. Label equations (1) and (2). Show ALL working."),
    ("4. Answer in Context",
     "Write a full sentence using the original words and units. Check your answer!"),
]
for num, hint in steps_ref:
    tbl_s = doc.add_table(rows=1, cols=2)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_s.columns[0].width = Inches(2)
    tbl_s.columns[1].width = Inches(4.8)
    c0 = tbl_s.cell(0, 0)
    set_cell_bg(c0, LGRAY)
    r0 = c0.paragraphs[0].add_run(num)
    r0.font.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = NAVY
    c1 = tbl_s.cell(0, 1)
    r1 = c1.paragraphs[0].add_run(hint)
    r1.font.size = Pt(10); r1.font.color.rgb = DGRAY
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 – GUIDED EXAMPLES
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, "📖  PART 2  |  Guided Examples  (follow along)", bg=GREEN)

# Example A
heading(doc, "Example A  —  Ticket Sales  ⭐", size=13, color=GREEN, space_before=8)
body(doc,
     "Adult tickets cost $12 and student tickets cost $7. "
     "A total of 350 tickets were sold, raising $3 150. "
     "How many adult and student tickets were sold?",
     size=11, color=DGRAY)

body(doc, "Step 1 — Define Variables:", bold=True, size=11, color=TEAL, space_before=6)
body(doc, "   Let a = number of adult tickets      Let s = number of student tickets",
     size=11, color=DGRAY)

body(doc, "Step 2 — Form Equations:", bold=True, size=11, color=GREEN, space_before=4)
body(doc, "   Total tickets:    a  +  s  =  350     … (1)", size=11, color=DGRAY)
body(doc, "   Total revenue:  12a + 7s  = 3 150   … (2)", size=11, color=DGRAY)

body(doc, "Step 3 — Solve (elimination):", bold=True, size=11, color=ORANGE, space_before=4)
body(doc, "   Multiply (1) by 7:   7a + 7s = 2 450  … (3)", size=11, color=DGRAY)
body(doc, "   Subtract (3) from (2):   5a = 700   ∴  a = 140", size=11, color=DGRAY)
body(doc, "   Substitute into (1):  140 + s = 350   ∴  s = 210", size=11, color=DGRAY)

body(doc, "Step 4 — Answer:", bold=True, size=11, color=PURPLE, space_before=4)
body(doc, "   140 adult tickets and 210 student tickets were sold.", size=11, color=DGRAY)
body(doc, "   Check: 140 + 210 = 350 ✓   and   12(140) + 7(210) = 1680 + 1470 = 3150 ✓",
     size=10, italic=True, color=DGRAY)
doc.add_paragraph()

# Example B
heading(doc, "Example B  —  Break-Even Analysis  ⭐⭐", size=13, color=GREEN, space_before=8)
body(doc,
     "Emma starts a cupcake business. Fixed costs are $240. Each cupcake costs $1.50 to make "
     "and sells for $4.50.",
     size=11, color=DGRAY)
body(doc, "(a)  Write equations for Cost (C) and Revenue (R) in terms of n (number of cupcakes).",
     size=11, bold=True, color=NAVY)
body(doc, "     C = _______________________________        R = _______________________________",
     size=11, color=DGRAY)
body(doc, "(b)  Find the break-even point (where C = R).", size=11, bold=True, color=NAVY,
     space_before=4)
blank_lines(doc, 4)
answer_box(doc, 2)

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 – FOUNDATION PRACTICE
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, "🌱  PART 3  |  Foundation Practice  (Questions 1–5)  ⭐  10 XP each", bg=GREEN)
body(doc, "Equations are partially set up for you. Complete the solution and answer in context.",
     size=10, italic=True, color=DGRAY)

foundation_qs = [
    ("Q1  —  Two Numbers  ⭐",
     "The sum of two numbers is 42 and their difference is 8. Find both numbers.",
     [("Let x = larger number,  y = smaller number", False),
      ("x  +  y  =  _______    … (1)        x  −  y  =  _______    … (2)", False),
      ("Solve by adding the equations:  2x = _______  ∴  x = _______", False),
      ("Substitute back:  y = _______", False)], 10),
    ("Q2  —  Fruit Shop  ⭐",
     "A bag of apples costs $a and a bag of oranges costs $r. "
     "3 bags of apples + 2 bags of oranges = $13. "
     "1 bag of apples + 4 bags of oranges = $11. Find the cost of each.",
     [("(1)  3a + 2r = _______       (2)  a + 4r = _______", False),
      ("Method chosen: ________________________________", False),
      ("Working:", False)], 12),
    ("Q3  —  Perimeter  ⭐",
     "A rectangle has perimeter 52 cm. Its length is 8 cm more than its width. "
     "Find the length and width.",
     [("Let l = length,  w = width", False),
      ("Perimeter equation: 2l + 2w = _______   →   l + w = _______    … (1)", False),
      ("Length/width relationship:  l − w = _______    … (2)", False),
      ("Solve:", False)], 12),
    ("Q4  —  Mixing Solutions  ⭐⭐",
     "A chemist mixes a 20% acid solution with a 50% acid solution to make 12 litres "
     "of a 35% acid solution. How many litres of each solution does he use?",
     [("Let x = litres of 20% solution,  y = litres of 50% solution", False),
      ("Total volume:  x + y = _______    … (1)", False),
      ("Acid equation:  0.2x + 0.5y = _______    … (2)", False),
      ("Working:", False)], 14),
    ("Q5  —  Mobile Plans  ⭐⭐",
     "Plan A: $25 per month + $0.10 per text. Plan B: $15 per month + $0.25 per text. "
     "Find the number of texts for which both plans cost the same.",
     [("Cost A:  C = _______  +  _______  × n", False),
      ("Cost B:  C = _______  +  _______  × n", False),
      ("Set equal and solve:  _______________________________", False)], 12),
]

for title, question, scaffold_lines, xp in foundation_qs:
    heading(doc, f"{title}  ({xp} XP)", size=12, color=GREEN, space_before=10)
    body(doc, question, size=11, color=DGRAY)
    doc.add_paragraph()
    for line_text, is_bold in scaffold_lines:
        body(doc, "   " + line_text, size=11, color=DGRAY, bold=is_bold, space_before=2)
    blank_lines(doc, 4)
    answer_box(doc, 2)

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 – STANDARD PRACTICE
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc,
    "📘  PART 4  |  Standard Practice  (Questions 6–9)  ⭐⭐  15 XP each", bg=TEAL)
body(doc, "Set up your own equations. Choose your method. Full working required.",
     size=10, italic=True, color=DGRAY)

standard_qs = [
    ("Q6  —  Age Problem  ⭐⭐",
     "Maria is three times as old as her daughter Lily. In 10 years, Maria will be twice Lily's age. "
     "Find their current ages.",
     16),
    ("Q7  —  Distance & Speed  ⭐⭐",
     "Two trains leave cities 480 km apart at the same time, travelling towards each other. "
     "Train A travels at 90 km/h and Train B at 70 km/h. "
     "When and where do they meet? (Hint: combined they cover 480 km together.)",
     18),
    ("Q8  —  Investment  ⭐⭐⭐",
     "Omar invests $8 000 in two accounts. Account X pays 4% annual interest, "
     "Account Y pays 6%. After one year, the total interest is $380. "
     "How much did Omar invest in each account?",
     18),
    ("Q9  —  Geometry  ⭐⭐⭐",
     "Two angles are supplementary (add to 180°). One angle is 24° more than three times the other. "
     "Find both angles. Then determine whether the lines with equations y = (first angle)x + 1 "
     "and y = (second angle)x − 3 are parallel, perpendicular, or neither.",
     20),
]

for title, question, xp in standard_qs:
    heading(doc, f"{title}  ({xp} XP)", size=12, color=TEAL, space_before=10)
    body(doc, question, size=11, color=DGRAY)
    body(doc, "Define variables:", bold=True, size=10, color=NAVY, space_before=4)
    body(doc, "   Let ___ = _______________        Let ___ = _______________",
         size=11, color=DGRAY)
    body(doc, "Equations:", bold=True, size=10, color=NAVY, space_before=2)
    body(doc, "   (1) ________________________________     (2) ________________________________",
         size=11, color=DGRAY)
    blank_lines(doc, 5)
    answer_box(doc, 2)

# ══════════════════════════════════════════════════════════════════════════════
# PART 6 – ADVANCED / REASONING
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc,
    "🔥  PART 5  |  Advanced & Reasoning  (Questions 10–14)  ⭐⭐⭐  20 XP each", bg=ORANGE)
body(doc, "Multi-step reasoning — show clear logical working. These test your depth of understanding.",
     size=10, italic=True, color=DGRAY)

advanced_qs = [
    ("Q10  —  Parallel Lines Analysis  ⭐⭐⭐",
     "For each pair of equations, determine (without solving) whether there is one solution, "
     "no solution, or infinite solutions. Justify your answer algebraically.\n\n"
     "(a)  3x + 2y = 12   and   6x + 4y = 24\n"
     "(b)  y = 4x − 3     and   2y = 8x + 1\n"
     "(c)  x + 2y = 7     and   2x − y = 4"),
    ("Q11  —  Break-Even (Extended)  ⭐⭐⭐",
     "A start-up makes wireless earbuds. Fixed costs: $12 000. Variable cost: $18 per pair. "
     "Selling price: $45 per pair.\n\n"
     "(a)  Write equations for Cost C and Revenue R in terms of n.\n"
     "(b)  Find the break-even point.\n"
     "(c)  How many pairs must they sell to make a profit of at least $5 400?"),
    ("Q12  —  Reverse Engineering  ⭐⭐⭐",
     "A pair of simultaneous equations has the solution x = 3, y = −2. "
     "Write TWO different pairs of equations that produce this solution. "
     "Explain how you constructed them."),
    ("Q13  —  Proof  ⭐⭐⭐⭐",
     "Prove algebraically that if two lines y = m₁x + c₁ and y = m₂x + c₂ are parallel "
     "(m₁ = m₂, c₁ ≠ c₂), then the system of equations has no solution."),
    ("Q14  —  Real-World Modelling  ⭐⭐⭐⭐",
     "Create your own real-world application problem involving simultaneous equations. "
     "Your problem must:\n"
     "  •  Be set in a realistic context\n"
     "  •  Have exactly one solution\n"
     "  •  Require both elimination and back-substitution\n"
     "  •  Include a full worked solution\n\n"
     "Exchange your problem with a partner and solve theirs!"),
]

for title, question in advanced_qs:
    heading(doc, title, size=12, color=ORANGE, space_before=10)
    body(doc, question, size=11, color=DGRAY)
    blank_lines(doc, 6)
    answer_box(doc, 2)

# ══════════════════════════════════════════════════════════════════════════════
# EXIT TICKET
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, "🪞  EXIT TICKET  |  (3 minutes)  —  Hand this in before you leave!", bg=PURPLE)

body(doc,
     "A jar contains 20-cent and 50-cent coins. There are 30 coins worth $12.00 in total. "
     "How many of each coin are there?",
     size=12, bold=True, color=NAVY)
blank_lines(doc, 5)
answer_box(doc, 2)

# Self-assessment
body(doc, "Self-Assessment  —  circle one:", bold=True, size=11, color=NAVY, space_before=8)
tbl_sa = doc.add_table(rows=1, cols=3)
tbl_sa.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, col) in enumerate([
    ("🟢 GREEN — I've got this!", GREEN),
    ("🟡 YELLOW — Nearly there…", RGBColor(0xC8, 0x96, 0x00)),
    ("🔴 RED — Need more practice", RED),
]):
    cell = tbl_sa.cell(0, i)
    set_cell_bg(cell, LGRAY)
    r = cell.paragraphs[0].add_run(label)
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = col

doc.add_paragraph()
body(doc, "📚  Homework: Exercise 1K — Working programs on the class portal", bold=True,
     size=11, color=NAVY, space_before=6)

doc.save("/home/user/mathsteaching/Ex 1K - Applications Worksheet.docx")
print("Worksheet saved successfully!")
